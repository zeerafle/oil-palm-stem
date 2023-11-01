from docx import Document
from docx.shared import Cm

import pandas as pd
from PIL import Image
import os

DATA_PATH = os.path.join('data', 'selected_dataset')
SAVE_PATH = os.path.join('data', 'rotated_test_data')
WIDTH_SIZE = 1.78 #cm

os.makedirs(os.path.join(SAVE_PATH, 'infected'), exist_ok=True)
os.makedirs(os.path.join(SAVE_PATH, 'normal'), exist_ok=True)

# Create a new Word document
doc = Document()

# Add a table with 401 rows and 3 columns
table = doc.add_table(rows=400, cols=3)

# read the true vs prediction result
df = pd.read_csv(r'result\true_pred_224_globalavgpool.csv')

# loop over the dataframe
for index, row in df.iterrows():
    class_name = row['y_true']
    if row['y_true'] == 'terinfeksi':
        class_name = 'infected'
    picture_path = os.path.join(DATA_PATH, class_name, row['file_name'])

    # rotate and save picture
    rotated_picture_path = os.path.join(SAVE_PATH, class_name, row['file_name'])
    image = Image.open(picture_path)
    image = image.rotate(-90, expand=True)
    image.save(rotated_picture_path)

    # set the value of the cell
    picture_cell = table.cell(index, 0)
    paragraph = picture_cell.paragraphs[0]
    # insert picture
    run = paragraph.add_run()
    run.add_picture(rotated_picture_path, width=Cm(WIDTH_SIZE))
    # insert y_test and y_pred text
    y_test_cell = table.cell(index, 1)
    y_pred_cell = table.cell(index, 2)
    y_test_cell.text = row['y_true']
    y_pred_cell.text = row['y_pred']

# Save the document
doc.save('table_with_picture.docx')